from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "scholarpath_release" / "ScholarPath_Full_Product_Team_Tracker_2026-08-07.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "17233A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
AMBER = "FFF3CD"
GREEN = "E8F3EC"
RED = "FCE8E6"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, bold=False, color=INK, size=10.2):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(cell, text, bold=False, color=INK, size=9.4):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    style_run(run, bold=bold, color=color, size=size)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, label in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        add_text(cell, label, bold=True, color=DARK, size=9.2)
    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row_values)):
            fill = None
            if status_col == index:
                if value == "Done":
                    fill = GREEN
                elif value in {"Partial", "In progress", "Ready to build"}:
                    fill = AMBER
                elif value == "Blocked":
                    fill = RED
                elif value == "Not started":
                    fill = LIGHT_GRAY
            if fill:
                shade(cell, fill)
            add_text(cell, value, bold=(status_col == index), size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=10.4)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = header.add_run("SCHOLARPATH  |  INTERNAL PRODUCT TRACKER")
    style_run(hrun, bold=True, color="6B7280", size=8.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run("Internal planning document - update after every verified change")
    style_run(frun, color="6B7280", size=8.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("ScholarPath full-product team tracker")
    style_run(run, bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Single operating document for Haidar, Codex, Claude, Gemini, product, design and engineering")
    style_run(run, color="5B6475", size=11.5)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    shade(callout.cell(0, 0), LIGHT_BLUE)
    p = callout.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run("Current decision: ")
    style_run(r, bold=True, color=DARK, size=10.5)
    r = p.add_run("the 14 August beta deadline is removed. We are building a full product. A module is only Done when it is implemented, connected, verified and safe to build upon.")
    style_run(r, color=INK, size=10.5)

    add_heading(doc, "How to use this document", 1)
    for bullet in [
        "Before starting work, claim a task in the priority ledger and add your name or agent name as owner.",
        "After work, update status, files changed, verification run and commit/PR reference in the Markdown mirror: SCHOLARPATH_TEAM_TRACKER.md.",
        "Do not describe UI-only work, demo data or unreviewed extracted data as complete.",
        "If two agents may edit the same module, one person coordinates the contract/schema first.",
    ]:
        add_bullet(doc, bullet)

    add_heading(doc, "Executive truth", 1)
    add_table(doc, ["Area", "Status", "What is true now", "What must happen next"], [
        ("Platform", "Partial", "Next.js shell, Supabase connection and protected admin shell exist; local build, typecheck and 24 tests pass.", "Prove live student/admin journeys and production deployment."),
        ("Opportunity data", "Partial", "Governed source registry, Edge worker, snapshots, candidates and review queue exist. EACEA has 220 discovery leads.", "Build official source packs, structured extractors and reviewer publishing."),
        ("Recommendations", "Partial", "Eligibility/explainability direction exists.", "Use reviewed published data only; add deterministic scoring contract and golden profiles."),
        ("User experience", "Partial", "Global green loader was replaced; admin navigation no longer mixes operations and admin.", "Founder mobile/desktop acceptance and interaction consistency."),
        ("Public release", "Partial", "Backend is active; local verification passed.", "Security hardening, Vercel, monitoring, backups and acceptance proof."),
    ], [1400, 800, 3500, 3660], status_col=1)

    add_heading(doc, "Module scoreboard", 1)
    modules = [
        ("App shell and navigation", "Partial", "Urbanist shell; neutral route progress; admin route cleanup.", "UX acceptance; remove remaining legacy visual inconsistency."),
        ("Authentication and roles", "Partial", "Supabase Auth and protected admin routes.", "Recovery and redirect QA; two-student and reviewer/admin isolation tests."),
        ("Student profile", "Partial", "Profile flow and validation structure.", "Live save proof; controlled dropdown data; completeness score."),
        ("Student pathway report", "Partial", "Report and pathway experience concept.", "Live sourced outputs with uncertainty, evidence impact and task links."),
        ("Recommendation engine", "Partial", "Deterministic eligibility and explainability direction.", "Versioned scoring contract; 30 golden profiles; fairness checks."),
        ("Tasks, deadlines and Kanban", "Partial", "Task/Kanban interface direction.", "Evidence-gap tasks, impact scoring, deadline reminders and completion sync."),
        ("Opportunity ingestion", "Partial", "Registry, worker, snapshots, candidates, review queue.", "Global source packs, parsers, structured fields, freshness monitoring."),
        ("Admin command center", "Partial", "Protected source, review and run tabs at /admin.", "Bulk review and source health/failure visibility."),
        ("Country intelligence", "Partial", "Data model and UX direction.", "Current sourced country/city facts: costs, visa, work, safety, community and salaries."),
        ("Institution directory", "Partial", "Database/UI support.", "Institutions, campuses, programmes, intakes, equivalence and document rules."),
        ("Rankings", "Not started", "Rankings intentionally are context, not recommendation logic.", "Permission/licensing decision, history and comparison UI."),
        ("Security and privacy", "Partial", "Protected APIs and RLS/auth structure.", "Advisor fixes, leaked-password protection, isolation proof, audit/retention review."),
        ("Deployment and observability", "Partial", "Local quality gates pass; Supabase connected.", "Vercel, monitoring, uptime, backup and rollback proof."),
    ]
    add_table(doc, ["Module", "Status", "Completed", "Remaining"], modules, [1550, 850, 3200, 3760], status_col=1)

    doc.add_page_break()
    add_heading(doc, "Opportunity coverage", 1)
    add_table(doc, ["Source area", "Status", "Current position", "Next step"], [
        ("EACEA / Erasmus Mundus", "Partial", "220 discovery candidates; never auto-publish.", "Adopt verified records, extract structured facts, reviewer publish."),
        ("Chevening, DAAD, Ireland, Netherlands", "Partial", "Official detail sources registered.", "Source-specific extractors and review queues."),
        ("Leeds, Saarland, Trinity", "Partial", "First university programme detail sources registered.", "Generalise into institution/course ingestion."),
        ("Commonwealth, Australia Awards, MEXT", "Ready to build", "Official sources researched.", "Adapter, parser and fixture for each source."),
        ("Turkiye, Stipendium Hungaricum, EduCanada", "Ready to build", "Official sources researched.", "Adapter, parser and fixture for each source."),
        ("Fulbright, Korea GKS", "Ready to build", "Official programme pages researched.", "Country-aware packs and eligibility normalisation."),
        ("Manaaki New Zealand, IsDB", "Partial", "Official portals identified.", "Confirm extraction policy then design parser."),
    ], [1800, 1100, 3150, 3310], status_col=1)

    add_heading(doc, "Priority task ledger", 1)
    add_table(doc, ["ID", "Priority", "Task", "Acceptance evidence"], [
        ("P0-01", "P0", "Publish versioned opportunity schema and mandatory review fields.", "Incomplete candidate cannot publish."),
        ("P0-02", "P0", "Add official source packs: Commonwealth, Australia Awards, MEXT, Turkiye, Stipendium, EduCanada, Fulbright and GKS.", "Each has host allowlist, parser fixture, schedule, source URL and failure reporting."),
        ("P0-03", "P0", "Complete candidate to reviewed/published/rejected reviewer workflow.", "Reviewer can publish a valid record and it appears in Discover."),
        ("P0-04", "P0", "Connect published opportunities to one deterministic recommendation contract.", "Discover, Today and Report show stable, explainable results."),
        ("P0-05", "P0", "Generate impact-prioritised Kanban tasks from missing eligibility evidence.", "Blocked recommendation links to owner, due date, impact and completion effect."),
        ("P0-06", "P0", "Run live student and admin browser acceptance journeys.", "Auth-profile-recommendation-report-task and source-review-publish both pass."),
        ("P0-07", "P0", "Resolve Supabase advisor issues and prove role isolation.", "No critical finding; unrelated students cannot access each other; staff roles proven."),
        ("P1-01", "P1", "Build country intelligence source/review workflow.", "Initial target countries have current, sourced facts with update dates."),
        ("P1-02", "P1", "Build institution, campus, programme and intake directory ingestion.", "Published directory data is distinct from opportunity data and filterable."),
        ("P1-03", "P1", "Add 30 golden profiles and recommendation regressions.", "Expected eligibility, reasons and ordering documented; tests pass."),
        ("P1-04", "P1", "Deploy Vercel with monitoring and rollback procedure.", "Production smoke test, alert destination and rollback documented."),
        ("P2-01", "P2", "Add rankings as transparent licensed context only.", "License recorded; rankings never override fit or affordability."),
        ("P2-02", "P2", "Complete mobile UX and interaction polish.", "Founder signs off desktop/mobile accessible states."),
    ], [700, 650, 4200, 3810])

    add_heading(doc, "Blockers that matter", 1)
    add_table(doc, ["Blocker", "Why it matters", "Unblock action"], [
        ("Few reviewed/published opportunities", "Global recommendations would be unreliable.", "Complete P0-01 to P0-03."),
        ("No recommendation regression set", "Scoring changes can silently harm students.", "Complete P1-03."),
        ("Security acceptance incomplete", "Public release risk.", "Complete P0-07 and verify live settings."),
        ("Vercel/monitoring not proven", "No dependable public release path.", "Complete P1-04."),
        ("Founder UX acceptance pending", "Quality remains subjective and inconsistent.", "Complete P0-06, then P2-02."),
    ], [2500, 3100, 3760])

    add_heading(doc, "Data safety rule", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ScholarPath may discover opportunities automatically, but it must only recommend records that are current, official-source-linked, structurally complete, reviewed and published. It must never invent scholarship, admission or visa probabilities.")
    style_run(r, bold=True, color=DARK, size=10.6)

    doc.core_properties.title = "ScholarPath Full Product Team Tracker"
    doc.core_properties.subject = "Shared product status and priority ledger"
    doc.core_properties.author = "ScholarPath"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
